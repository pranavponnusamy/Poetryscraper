import requests
from bs4 import BeautifulSoup
from html import escape

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from googlesearch import search

def manual_get_poem_urls(poem_names=['to helen']):
    
    for poem in poem_names:
        url = f"https://www.google.com/search?q={poem}"
        response = requests.get(url)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')
        result_divs = soup.find(id = "rcnt")
        print(result_divs)


def get_poem_names():
    names = []
    f = open("src/poem-names-authors.txt", "rb")
    for name in f:
        names.append((name.replace(b'\"', b'').strip().decode("utf-8")))
    f.close()
    print(names)
    return names
        

def generate_poem_urls(poem_names=['to helen']):
    fileout = "urls.txt"
    output = open(fileout,'w')
    count = 0
    for poem in poem_names:
        query = f"{poem} poem"
        link_exists = False
        
        for j in search(query, tld="co.in", num=5, stop=5, pause=2):
            if ("poetryfoundation" in j)  and ("poem" in j) and (link_exists == False):
                    print(j)
                    print(j, file=output)
                    link_exists = True
            
           
        if not link_exists:                
            for j in search(query+" allpoetry", tld="co.in", num=3, stop=3, pause=2):  
                if ("allpoetry" in j)  and (link_exists == False):
                    for x in range(0, len(poem.split())-1):
                        word = poem.split()
                        if (word[x].lower() in j.lower()) and link_exists == False:
                            print(j)
                            print(j, file=output)
                            link_exists = True      
                
            
        if not link_exists:
            print(f"{poem} not found")
            print(f"{poem} not found", file=output)
        print(count)
        count += 1


def scape_poem_allPoetry(poem_urls=["https://allpoetry.com/poem/8448403-The-Raven-by-Edgar-Allan-Poe"]):
    doc = docx.Document()
    
    for url in poem_urls:  
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, "html.parser")
        
        # Find the relevant section of the poem
        poem_div = soup.find("div", class_="poem_body")
    
        if poem_div:            
            poem_name = soup.find("h1", class_="title vcard item otitle_8448403").text
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Center Alignment
            title_run = title_paragraph.add_run(poem_name)
            title_run.bold = True
            title_run.font.size = Pt(18)

            stanza_chunk = poem_div.find("div", class_="orig_8448403").find("p").text
            stanzas = doc.add_paragraph()
            stanzas.alignment = WD_ALIGN_PARAGRAPH.LEFT
            stanzas.add_run(stanza_chunk).font.size = Pt(11)

            doc.add_page_break()

    doc.save("poems.docx")
        
        

def scrape_poem_pfoundation(poem_urls=[]):
    doc = docx.Document()

    # fileout = "poems.txt"
    # output = open(fileout, "w")

    url = "https://www.poetryfoundation.org/poems/44888/to-helen"

    for url in list:
        try:
            response = requests.get(url)
            response.raise_for_status()  # Raise an exception for bad status codes (like 404)
            soup = BeautifulSoup(response.content, "html.parser")

            # Find the relevant section of the poem
            poem_div = soup.find("div", class_="o-poem")

            if poem_div:
                title = unescape(soup.find("h1").text).strip()
                author = unescape(
                    soup.find("span", class_="c-txt c-txt_attribution").text
                ).strip()
                print(author)
                lines = []
                for h1 in poem_div.find_all("div"):
                    lines.append(unescape(h1.text))

                # print(f"Title: {title}", file=output)
                # print("-" * len(title), file=output)  # Separator

                title_paragraph = doc.add_paragraph()
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Center Alignment
                title_run = title_paragraph.add_run(title)
                title_run.bold = True
                title_run.font.size = Pt(18)

                author_paragraph = doc.add_paragraph()
                author_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Center Alignment
                author_paragraph.add_run(author + "\n").font.size = Pt(10)

                for line in lines:
                    # print(line, file=output)
                    stanza = doc.add_paragraph()
                    stanza.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    stanza.add_run(line).font.size = Pt(11)

            else:
                print("Poem not found on the page.")

        except requests.exceptions.RequestException as e:
            print(f"An error occurred during the request: {e}")

    doc.save("poems.docx")


if __name__ == "__main__":
    # scrape_poem()
    generate_poem_urls(get_poem_names())
    # get_poem_names()
    # manual_get_poem_urls()
    # scape_poem_allPoetry()


